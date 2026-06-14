# test_module2.py
"""
Test script cho Module 2: Text Extractor

Test cases:
1. Extract text từ PDF thật
2. Extract text từ DOCX thật
3. Test PDF rỗng (không có text)
"""

import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent))

from modules.input_handler import validate_and_process
from modules.text_extractor import extract_text, ExtractionError


def create_test_files():
    """Tạo sample PDF và DOCX để test"""
    test_dir = Path(__file__).parent / "test_files"
    test_dir.mkdir(exist_ok=True)

    # Tạo PDF thật dùng reportlab (nếu có) hoặc dùng PDF có sẵn
    # Thử tạo DOCX thật
    try:
        from docx import Document

        # DOCX 1: Contract
        doc = Document()
        doc.add_heading('SERVICE CONTRACT', 0)
        doc.add_paragraph('This agreement is made between AvePoint Inc and John Doe.')
        doc.add_paragraph('TERMS AND CONDITIONS:')
        doc.add_paragraph('1. The service will commence on January 1, 2024.')
        doc.add_paragraph('2. Payment of $5,000 is due within 30 days.')
        doc.add_paragraph('3. Both parties agree to the terms hereby stated.')
        doc.add_paragraph('Signed: John Doe | Date: 2024-01-15')
        doc.save(str(test_dir / "contract_test.docx"))
        print("✓ Tạo contract_test.docx thành công")

        # DOCX 2: Invoice
        doc2 = Document()
        doc2.add_heading('INVOICE #INV-2024-001', 0)
        doc2.add_paragraph('From: AvePoint Inc')
        doc2.add_paragraph('To: John Doe')
        doc2.add_paragraph('Date: 2024-01-15')
        doc2.add_paragraph('Amount Due: $1,500.00')
        doc2.add_paragraph('Description: Software license fee')
        doc2.add_paragraph('Payment due by: 2024-02-15')
        doc2.save(str(test_dir / "invoice_test.docx"))
        print("✓ Tạo invoice_test.docx thành công")

    except Exception as e:
        print(f"⚠ Không tạo được DOCX: {e}")

    # Tạo PDF thật
    try:
        # Tạo PDF đơn giản theo chuẩn
        pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj

2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj

3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj

4 0 obj
<< /Length 200 >>
stream
BT
/F1 12 Tf
50 750 Td
(INVOICE DOCUMENT) Tj
0 -20 Td
(Invoice Number: INV-001) Tj
0 -20 Td
(Amount: $500.00) Tj
0 -20 Td
(Due Date: 2024-02-01) Tj
0 -20 Td
(From: Company ABC) Tj
ET
endstream
endobj

5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj

xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000516 00000 n

trailer
<< /Size 6 /Root 1 0 R >>
startxref
600
%%EOF"""

        with open(str(test_dir / "invoice_test.pdf"), 'wb') as f:
            f.write(pdf_content)
        print("✓ Tạo invoice_test.pdf thành công")

    except Exception as e:
        print(f"⚠ Không tạo được PDF: {e}")

    return test_dir


def run_extraction_test(label: str, file_path: Path):
    """Chạy 1 test case và in kết quả"""
    print(f"\n{'='*60}")
    print(f"TEST: {label}")
    print(f"{'='*60}")
    print(f"File: {file_path.name}")

    try:
        # Bước 1: Module 1 validate
        file_obj = validate_and_process(str(file_path))
        print(f"✓ Module 1 OK: {file_obj}")

        # Bước 2: Module 2 extract
        result = extract_text(file_obj)
        print(f"✓ Module 2 OK: {result}")

        if result.raw_text:
            print(f"\n📄 TEXT EXTRACTED (first 200 chars):")
            print("-" * 40)
            print(result.raw_text[:200])
            print("-" * 40)
        else:
            print(f"⚠ Không có text: {result.error}")

    except Exception as e:
        print(f"❌ Lỗi: {type(e).__name__}: {str(e)}")


def main():
    print("=" * 60)
    print("MODULE 2 TEST SUITE: Text Extractor")
    print("=" * 60)

    print("\n📁 Creating test files...")
    test_dir = create_test_files()

    # Test DOCX files
    if (test_dir / "contract_test.docx").exists():
        run_extraction_test("Extract text từ Contract DOCX",
                           test_dir / "contract_test.docx")

    if (test_dir / "invoice_test.docx").exists():
        run_extraction_test("Extract text từ Invoice DOCX",
                           test_dir / "invoice_test.docx")

    # Test PDF file
    if (test_dir / "invoice_test.pdf").exists():
        run_extraction_test("Extract text từ Invoice PDF",
                           test_dir / "invoice_test.pdf")

    print(f"\n{'='*60}")
    print("TEST SUITE COMPLETED")
    print(f"{'='*60}")
    print("\n💡 Next step: Giải thích lại cho Study Agent trước khi sang Module 3!")


if __name__ == "__main__":
    main()
